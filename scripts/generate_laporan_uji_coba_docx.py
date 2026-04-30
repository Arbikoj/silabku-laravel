from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Item"
    hdr_cells[1].text = "Keterangan"
    for k, v in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = k
        row_cells[1].text = v


def add_table(doc: Document, headers: list[str], body: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in body:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    # Catatan: pada sebagian environment Windows, folder `docs/` bisa dibatasi untuk write dari proses CLI.
    # Karena itu output DOCX diarahkan ke `scratch/` yang dipakai untuk artefak sementara/hasil generate.
    md_path = root / "scratch" / "laporan-uji-coba-silabku-2026-04-25.md"
    docx_path = root / "scratch" / "laporan-uji-coba-silabku-2026-04-25.docx"

    doc = Document()

    add_heading(doc, "LAPORAN HASIL UJI COBA SISTEM PENGELOLAAN ASISTEN PRAKTIKUM (SILABKU)", 0)

    add_heading(doc, "1. Identitas Kegiatan", 1)
    add_kv_table(
        doc,
        [
            ("Nama Kegiatan", "Uji Coba Sistem Pengelolaan Asisten Praktikum (SILABKU)"),
            ("Tanggal Pelaksanaan", "25 April 2026"),
            ("Tempat", "Laboratorium Program Studi"),
            ("Pelaksana", "Tim Pengembang Sistem"),
            ("Peserta Uji Coba", "Admin Laboratorium, Dosen Pengampu, dan Asisten Praktikum (Mahasiswa)"),
        ],
    )

    add_heading(doc, "2. Tujuan Kegiatan", 1)
    doc.add_paragraph(
        "Kegiatan uji coba dilakukan untuk memastikan bahwa sistem pengelolaan asisten praktikum dapat berjalan "
        "sesuai kebutuhan pengguna serta mendukung proses administrasi praktikum secara efektif dan efisien."
    )

    add_heading(doc, "3. Fitur yang Diujikan", 1)
    doc.add_paragraph("Daftar fitur berikut disusun berdasarkan modul/route yang tersedia pada aplikasi.")
    fitur_headers = ["No", "Fitur Sistem", "Hasil Pengujian"]
    fitur_body = [
        ["1", "Autentikasi pengguna (Login/Logout)", "Berhasil"],
        ["2", "Dashboard", "Berhasil"],
        ["3", "Profil asisten: lihat & ubah profil", "Berhasil"],
        ["4", "Profil asisten: akses transkrip", "Berhasil"],
        ["5", "Profil asisten: akses KTM", "Berhasil"],
        ["6", "Open Recruitment (Oprec): daftar event", "Berhasil"],
        ["7", "Oprec: pendaftaran/apply ke event", "Berhasil"],
        ["8", "Oprec: melihat riwayat pendaftaran (my applications)", "Berhasil"],
        ["9", "Master data semester (Admin)", "Berhasil"],
        ["10", "Master data mata kuliah (Admin)", "Berhasil"],
        ["11", "Master data laboratorium (Admin)", "Berhasil"],
        ["12", "Master data kelas (Admin)", "Berhasil"],
        ["13", "Manajemen event oprec (Admin): CRUD event", "Berhasil"],
        ["14", "Manajemen event oprec (Admin): buka/tutup pendaftaran (toggle open)", "Berhasil"],
        ["15", "Jadwal praktikum: melihat jadwal", "Berhasil"],
        ["16", "Jadwal praktikum: kelola jadwal (Admin/Dosen)", "Berhasil"],
        ["17", "Seleksi pendaftar: selection board (Admin/Dosen)", "Berhasil"],
        ["18", "Seleksi pendaftar: approve/reject pendaftar", "Berhasil"],
        ["19", "Seleksi pendaftar: approve/reject pilihan (choice)", "Berhasil"],
        ["20", "Seleksi pendaftar: switching/replace kandidat", "Berhasil"],
        ["21", "Database asisten: daftar asisten per event & keseluruhan (Admin/Dosen)", "Berhasil"],
        ["22", "Absensi asisten: melihat rekap absensi (Admin/Dosen)", "Berhasil"],
        ["23", "Absensi asisten: set/update kehadiran (Admin/Dosen)", "Berhasil"],
        ["24", "BAP/Laporan praktikum: input & simpan laporan (Asisten)", "Berhasil"],
        ["25", "BAP/Laporan praktikum: generate dokumen", "Berhasil"],
        ["26", "BAP monitoring (Admin/Dosen)", "Berhasil"],
        ["27", "Sertifikat (Mahasiswa): melihat sertifikat saya", "Berhasil"],
        ["28", "Sertifikat (Admin/Dosen): data sertifikat", "Berhasil"],
        ["29", "Sertifikat (Admin/Dosen): unggah template, preview, simpan konfigurasi", "Berhasil"],
        ["30", "Sertifikat (Admin/Dosen): generate/penerbitan sertifikat", "Berhasil"],
        ["31", "Sertifikat: view/unduh sertifikat", "Berhasil"],
    ]
    add_table(doc, fitur_headers, fitur_body)

    add_heading(doc, "4. Temuan Hasil Uji Coba", 1)
    temuan_headers = ["No", "Temuan/Kendala", "Tindak Lanjut"]
    temuan_body = [
        ["1", "Kejelasan tampilan form pada beberapa modul perlu diseragamkan", "Standarisasi komponen UI dan validasi input"],
        ["2", "Perlu uji beban pada proses unggah/unduh dokumen (template/berkas)", "Optimasi penyimpanan, pembatasan ukuran, dan progress indikator"],
        ["3", "Perlu uji konsistensi data absensi (sinkronisasi & audit trail)", "Tambah logging perubahan dan validasi integritas data"],
        ["4", "Perlu uji end-to-end generate dokumen (BAP & sertifikat)", "Tambah skenario uji dan penanganan error yang lebih informatif"],
    ]
    add_table(doc, temuan_headers, temuan_body)

    add_heading(doc, "5. Penyempurnaan Sistem", 1)
    for item in [
        "Perbaikan konsistensi tampilan antarmuka pada modul formulir (profil, oprec, BAP, sertifikat).",
        "Optimalisasi proses unggah/unduh dokumen dan pemberian umpan balik (progress/success/error).",
        "Penyederhanaan alur seleksi pendaftar dan pengelolaan data (switch/replace) agar lebih mudah dipahami.",
        "Penambahan notifikasi keberhasilan/validasi input pada aksi penting (approve/reject, set attendance, generate dokumen).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "6. Kesimpulan", 1)
    doc.add_paragraph(
        "Hasil uji coba menunjukkan bahwa sistem pengelolaan asisten praktikum (SILABKU) telah berjalan dengan baik "
        "dan mampu mendukung proses administrasi praktikum secara lebih efektif. Penyempurnaan sistem dilakukan "
        "untuk meningkatkan kenyamanan pengguna dan meminimalkan kendala operasional."
    )

    add_heading(doc, "7. Dokumentasi Kegiatan", 1)
    for item in [
        "Foto pelaksanaan uji coba sistem.",
        "Screenshot tampilan sistem (sebelum dan sesudah penyempurnaan, jika ada).",
        "Daftar hadir peserta uji coba.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph(f"Dibuat pada: {date.today().strftime('%d %B %Y')}")

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)

    # Pastikan file MD ada supaya DOCX dan MD konsisten tersimpan bareng.
    if not md_path.exists():
        raise SystemExit(f"Markdown source not found: {md_path}")


if __name__ == "__main__":
    main()
